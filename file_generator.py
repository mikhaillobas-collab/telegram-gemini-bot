import io
import json
import logging
import os
import re
import sys
import tempfile
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

import visual_generator

logger = logging.getLogger(__name__)

# --- Настройка шрифтов для ReportLab (Cyrillic PDF) ---
CYRILLIC_FONT = "Helvetica"
CYRILLIC_BOLD = "Helvetica-Bold"

FONT_CANDIDATES = [
    # Linux (Debian / Ubuntu / Docker / bothost)
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ("/usr/share/fonts/dejavu/DejaVuSans.ttf", "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
    ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
    # Windows
    ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
    ("C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/calibrib.ttf"),
]

for reg_path, bold_path in FONT_CANDIDATES:
    if os.path.exists(reg_path):
        try:
            pdfmetrics.registerFont(TTFont("CustomCyrillic", reg_path))
            CYRILLIC_FONT = "CustomCyrillic"
            if os.path.exists(bold_path):
                pdfmetrics.registerFont(TTFont("CustomCyrillicBold", bold_path))
                CYRILLIC_BOLD = "CustomCyrillicBold"
            else:
                CYRILLIC_BOLD = "CustomCyrillic"
            break
        except Exception as e:
            logger.warning("Не удалось зарегистрировать шрифт %s: %s", reg_path, e)


# ==============================================================================
# 1. Генератор Excel (.xlsx)
# ==============================================================================

def create_excel_file(data: Dict[str, Any]) -> Tuple[bytes, str]:
    filename = data.get("filename", "Таблица.xlsx")
    if not filename.endswith(".xlsx"):
        filename += ".xlsx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = str(data.get("sheet", "Лист 1"))[:30]

    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    title_font = Font(name="Calibri", size=14, bold=True, color="1F4E78")
    data_font = Font(name="Calibri", size=11)
    
    thin_border = Border(
        left=Side(style="thin", color="D9D9D9"),
        right=Side(style="thin", color="D9D9D9"),
        top=Side(style="thin", color="D9D9D9"),
        bottom=Side(style="thin", color="D9D9D9"),
    )
    zebra_fill = PatternFill(start_color="F2F5F9", end_color="F2F5F9", fill_type="solid")

    current_row = 1

    title = data.get("title")
    if title:
        ws.cell(row=current_row, column=1, value=str(title))
        ws.cell(row=current_row, column=1).font = title_font
        current_row += 2

    headers = data.get("headers", [])
    rows = data.get("rows", [])

    if headers:
        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=str(header))
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border
        current_row += 1

    for r_idx, row in enumerate(rows):
        is_even = (r_idx % 2 == 1)
        for c_idx, val in enumerate(row, start=1):
            converted_val = val
            if isinstance(val, str):
                val_clean = val.replace(" ", "").replace(",", ".")
                try:
                    if "." in val_clean:
                        converted_val = float(val_clean)
                    else:
                        converted_val = int(val_clean)
                except ValueError:
                    converted_val = val

            cell = ws.cell(row=current_row, column=c_idx, value=converted_val)
            cell.font = data_font
            cell.border = thin_border
            if is_even:
                cell.fill = zebra_fill

            if isinstance(converted_val, (int, float)):
                cell.alignment = Alignment(horizontal="right", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center")

        current_row += 1

    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or "")
            max_len = max(max_len, len(val_str))
        ws.column_dimensions[col_letter].width = min(max(max_len + 4, 12), 50)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue(), filename


# ==============================================================================
# 2. Генератор Word (.docx) со встроенными таблицами и графиками
# ==============================================================================

def create_word_file(data: Dict[str, Any]) -> Tuple[bytes, str]:
    """
    Генерирует расширенный документ Word (.docx).
    Может содержать разделы, маркированные списки, таблицы и встроенные графики/схемы.
    """
    filename = data.get("filename", "Документ.docx")
    if not filename.endswith(".docx"):
        filename += ".docx"

    doc = docx.Document()

    # Заголовок документа
    title_text = data.get("title")
    if title_text:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(str(title_text))
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 120)

    subtitle_text = data.get("subtitle")
    if subtitle_text:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(str(subtitle_text))
        sub_run.font.name = "Calibri"
        sub_run.font.size = Pt(11)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(120, 120, 120)

    # 1. Если в документе запрошена встроенная схема родства (scheme)
    if "scheme" in data and isinstance(data["scheme"], dict):
        try:
            scheme_bytes, _ = visual_generator.create_kinship_tree(data["scheme"])
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            doc.add_picture(io.BytesIO(scheme_bytes), width=Inches(6.2))
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
        except Exception as e:
            logger.exception("Ошибка вставки схемы в Word: %s", e)

    # 2. Если в документе запрошен встроенный график (chart)
    if "chart" in data and isinstance(data["chart"], dict):
        try:
            chart_bytes, _ = visual_generator.create_chart(data["chart"])
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.8))
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
        except Exception as e:
            logger.exception("Ошибка вставки графика в Word: %s", e)

    # Разделы документа
    sections = data.get("sections", [])
    for sec in sections:
        heading = sec.get("heading")
        if heading:
            h = doc.add_heading(level=2)
            h_run = h.add_run(str(heading))
            h_run.font.name = "Calibri"
            h_run.font.color.rgb = RGBColor(31, 78, 120)

        for p_text in sec.get("paragraphs", []):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(str(p_text))
            run.font.name = "Calibri"
            run.font.size = Pt(11)

        for item in sec.get("bullet_points", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(item))
            run.font.name = "Calibri"
            run.font.size = Pt(11)

        # Секция может содержать свою локальную таблицу
        if "table" in sec and isinstance(sec["table"], dict):
            t_data = sec["table"]
            headers = t_data.get("headers", [])
            rows = t_data.get("rows", [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Light Shading Accent 1"
                for col_idx, h in enumerate(headers):
                    table.cell(0, col_idx).text = str(h)
                for row_idx, row in enumerate(rows, start=1):
                    for col_idx, val in enumerate(row):
                        table.cell(row_idx, col_idx).text = str(val)
                doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Общая таблица внутри Word документа (если указана на верхнем уровне)
    table_data = data.get("table")
    if table_data and "headers" in table_data:
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if headers:
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Shading Accent 1"
            for col_idx, h in enumerate(headers):
                table.cell(0, col_idx).text = str(h)
            for row_idx, row in enumerate(rows, start=1):
                for col_idx, val in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(val)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), filename


# ==============================================================================
# 3. Генератор PDF-презентаций (Альбомный формат A4)
# ==============================================================================

def create_pdf_presentation(data: Dict[str, Any]) -> Tuple[bytes, str]:
    filename = data.get("filename", "Презентация.pdf")
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    buf = io.BytesIO()
    page_w, page_h = landscape(A4)
    c = canvas.Canvas(buf, pagesize=(page_w, page_h))

    slides = data.get("slides", [])
    title = data.get("title", "Презентация")
    subtitle = data.get("subtitle", "")
    total_slides = max(1 + len(slides), 1)

    # Слайд 1: Титульный
    c.setFillColor(colors.HexColor("#1A365D"))
    c.rect(0, 0, page_w, page_h, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#3182CE"))
    c.rect(0, page_h - 15, page_w, 15, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#63B3ED"))
    c.rect(0, 0, page_w, 10, fill=1, stroke=0)

    c.setFillColor(colors.white)
    c.setFont(CYRILLIC_BOLD, 30)
    c.drawCentredString(page_w / 2, page_h / 2 + 30, title)

    if subtitle:
        c.setFillColor(colors.HexColor("#E2E8F0"))
        c.setFont(CYRILLIC_FONT, 16)
        c.drawCentredString(page_w / 2, page_h / 2 - 25, subtitle)

    c.showPage()

    # Слайды контента
    for s_idx, slide in enumerate(slides, start=2):
        s_title = slide.get("title", f"Слайд {s_idx - 1}")
        points = slide.get("points", [])

        c.setFillColor(colors.HexColor("#F8FAFC"))
        c.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#1E3A8A"))
        c.rect(0, page_h - 75, page_w, 75, fill=1, stroke=0)

        c.setFillColor(colors.white)
        c.setFont(CYRILLIC_BOLD, 20)
        c.drawString(40, page_h - 48, s_title)

        c.setFillColor(colors.HexColor("#64748B"))
        c.setFont(CYRILLIC_FONT, 10)
        c.drawRightString(page_w - 40, 22, f"Слайд {s_idx} из {total_slides}")

        c.setStrokeColor(colors.HexColor("#CBD5E1"))
        c.setLineWidth(1)
        c.line(40, 42, page_w - 40, 42)

        current_y = page_h - 120
        line_height = 36

        c.setFont(CYRILLIC_FONT, 14)
        c.setFillColor(colors.HexColor("#1E293B"))

        for pt in points:
            if current_y < 65:
                break
            c.setFillColor(colors.HexColor("#2563EB"))
            c.circle(55, current_y + 5, 4, fill=1, stroke=0)

            c.setFillColor(colors.HexColor("#1E293B"))
            pt_str = str(pt)
            words = pt_str.split()
            current_line = []
            for w in words:
                test_line = " ".join(current_line + [w])
                if c.stringWidth(test_line, CYRILLIC_FONT, 14) < (page_w - 120):
                    current_line.append(w)
                else:
                    c.drawString(75, current_y, " ".join(current_line))
                    current_y -= 20
                    current_line = [w]

            if current_line:
                c.drawString(75, current_y, " ".join(current_line))
                current_y -= line_height

        c.showPage()

    c.save()
    buf.seek(0)
    return buf.getvalue(), filename


# ==============================================================================
# 4. Генератор многостраничных отчетов в PDF (Формат А4 вертикальный)
# ==============================================================================

def create_pdf_report(data: Dict[str, Any]) -> Tuple[bytes, str]:
    """
    Генерирует представительский PDF-отчет (формат A4) со стилями, таблицами и графиками.
    """
    filename = data.get("filename", "Отчет.pdf")
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    # Стили с поддержкой кириллицы
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontName=CYRILLIC_BOLD,
        fontSize=20,
        leading=26,
        textColor=colors.HexColor("#1E3A8A"),
        spaceAfter=8,
        alignment=1
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        parent=styles["Normal"],
        fontName=CYRILLIC_FONT,
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=18,
        alignment=1
    )
    heading_style = ParagraphStyle(
        "ReportHeading2",
        parent=styles["Heading2"],
        fontName=CYRILLIC_BOLD,
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#1E40AF"),
        spaceBefore=14,
        spaceAfter=8
    )
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["Normal"],
        fontName=CYRILLIC_FONT,
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#1E293B"),
        spaceAfter=6
    )

    story = []

    # Заголовок
    if data.get("title"):
        story.append(Paragraph(data["title"], title_style))
    if data.get("subtitle"):
        story.append(Paragraph(data["subtitle"], subtitle_style))

    # Если передан график для отчета
    if "chart" in data and isinstance(data["chart"], dict):
        try:
            chart_bytes, _ = visual_generator.create_chart(data["chart"])
            story.append(RLImage(io.BytesIO(chart_bytes), width=480, height=280))
            story.append(Spacer(1, 15))
        except Exception as e:
            logger.exception("Ошибка вставки графика в PDF-отчет: %s", e)

    # Если передана схема для отчета
    if "scheme" in data and isinstance(data["scheme"], dict):
        try:
            scheme_bytes, _ = visual_generator.create_kinship_tree(data["scheme"])
            story.append(RLImage(io.BytesIO(scheme_bytes), width=490, height=310))
            story.append(Spacer(1, 15))
        except Exception as e:
            logger.exception("Ошибка вставки схемы в PDF-отчет: %s", e)

    # Разделы отчета
    sections = data.get("sections", [])
    for sec in sections:
        if sec.get("heading"):
            story.append(Paragraph(sec["heading"], heading_style))

        for p_text in sec.get("paragraphs", []):
            story.append(Paragraph(p_text, body_style))

        # Таблица в разделе
        t_data = sec.get("table")
        if t_data and isinstance(t_data, dict):
            headers = [Paragraph(f"<b>{h}</b>", body_style) for h in t_data.get("headers", [])]
            table_rows = [headers]
            for row in t_data.get("rows", []):
                table_rows.append([Paragraph(str(val), body_style) for val in row])

            if len(table_rows) > 1:
                t = Table(table_rows, hAlign="CENTER")
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ]))
                story.append(Spacer(1, 8))
                story.append(t)
                story.append(Spacer(1, 12))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue(), filename


# ==============================================================================
# 5. Песочница исполнения Python-кода (Code Runner)
# ==============================================================================

def run_python_sandbox(code: str) -> Tuple[str, List[Tuple[bytes, str]]]:
    """
    Безопасно выполняет сгенерированный код в изолированной директории.
    Перехватывает созданные файлы (.png, .xlsx, .docx, .pdf) и возвращает их.
    """
    output_files: List[Tuple[bytes, str]] = []
    stdout_buf = io.StringIO()

    with tempfile.TemporaryDirectory() as tmp_dir:
        old_cwd = os.getcwd()
        old_stdout = sys.stdout
        try:
            os.chdir(tmp_dir)
            sys.stdout = stdout_buf

            # Пространство имен с готовыми библиотеками
            exec_globals = {
                "__builtins__": __builtins__,
                "plt": visual_generator.plt,
                "sns": visual_generator.sns,
                "docx": docx,
                "openpyxl": openpyxl,
                "io": io,
                "os": os,
            }
            exec(code, exec_globals)

            # Собираем все созданные файлы
            for fname in os.listdir(tmp_dir):
                fpath = os.path.join(tmp_dir, fname)
                if os.path.isfile(fpath):
                    with open(fpath, "rb") as f:
                        output_files.append((f.read(), fname))

        except Exception as e:
            logger.exception("Ошибка при исполнении Python-кода: %s", e)
            return f"❌ Ошибка выполнения кода: {e}", []
        finally:
            os.chdir(old_cwd)
            sys.stdout = old_stdout

    console_output = stdout_buf.getvalue().strip()
    return console_output, output_files


# ==============================================================================
# 6. Главный парсер блоков и маршрутизатор генерации
# ==============================================================================

EXCEL_PATTERN = re.compile(r"```(?:excel:data|excel|table:data)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
DOC_PATTERN = re.compile(r"```(?:doc:data|word:data|docx:data|doc)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
PRESENTATION_PATTERN = re.compile(r"```(?:presentation:data|pdf:presentation|slides:data)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
CHART_PATTERN = re.compile(r"```(?:chart:data|chart|graph:data)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
SCHEME_PATTERN = re.compile(r"```(?:scheme:data|tree:data|kinship:data|diagram:data)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
FLOWCHART_PATTERN = re.compile(r"```(?:flowchart:data|flow:data|process:data)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
REPORT_PATTERN = re.compile(r"```(?:report:data|pdf:report|document:data)\s*(\{.*?\})\s*```", re.DOTALL | re.IGNORECASE)
PYTHON_EXEC_PATTERN = re.compile(r"```(?:python:exec|python:file)\s*(.*?)\s*```", re.DOTALL | re.IGNORECASE)


def extract_and_generate_files(raw_text: str) -> Tuple[str, List[Tuple[bytes, str]]]:
    """
    Анализирует ответ модели, находит блоки данных и генерирует файлы.
    """
    cleaned_text = raw_text
    files: List[Tuple[bytes, str]] = []

    # 1. Графики (Chart)
    for match in CHART_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = visual_generator.create_chart(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока Chart: %s", e)

    # 2. Схемы родства и генеалогические деревья (Scheme/Tree)
    for match in SCHEME_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = visual_generator.create_kinship_tree(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока Scheme: %s", e)

    # 3. Блок-схемы процессов (Flowchart)
    for match in FLOWCHART_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = visual_generator.create_flowchart(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока Flowchart: %s", e)

    # 4. Таблицы Excel (.xlsx)
    for match in EXCEL_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = create_excel_file(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока Excel: %s", e)

    # 5. Документы Word (.docx) со встроенными таблицами/графиками
    for match in DOC_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = create_word_file(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока Word: %s", e)

    # 6. PDF-отчеты (Report A4)
    for match in REPORT_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = create_pdf_report(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока PDF Report: %s", e)

    # 7. PDF-презентации
    for match in PRESENTATION_PATTERN.finditer(raw_text):
        try:
            data = json.loads(match.group(1).strip())
            file_bytes, filename = create_pdf_presentation(data)
            files.append((file_bytes, filename))
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка парсинга блока Презентации: %s", e)

    # 8. Исполнение Python-кода (Code Runner)
    for match in PYTHON_EXEC_PATTERN.finditer(raw_text):
        try:
            code = match.group(1).strip()
            _, exec_files = run_python_sandbox(code)
            files.extend(exec_files)
            cleaned_text = cleaned_text.replace(match.group(0), "").strip()
        except Exception as e:
            logger.warning("Ошибка исполнения блока Python: %s", e)

    return cleaned_text, files
