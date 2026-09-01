import base64
import io
import logging
from typing import Optional, Tuple

import docx
from pypdf import PdfReader

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {
    ".txt", ".csv", ".json", ".md", ".py", ".js", ".ts", ".html", ".css",
    ".xml", ".yaml", ".yml", ".log", ".sql", ".sh", ".bat", ".ini", ".cfg",
    ".env", ".java", ".c", ".cpp", ".cs", ".go", ".rs", ".php", ".rb"
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}


def process_file_bytes(file_bytes: bytes, filename: str, mime_type: Optional[str] = None) -> Tuple[str, bool]:
    """
    Обрабатывает байты файла.
    Возвращает кортеж: (содержимое, is_image: bool).
    Если is_image == True, содержимое — это base64 строка.
    Если is_image == False, содержимое — это извлечённый текст.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # 1. Проверка на изображение
    if ext in IMAGE_EXTENSIONS or (mime_type and mime_type.startswith("image/")):
        b64 = base64.b64encode(file_bytes).decode("utf-8")
        return b64, True

    # 2. PDF файл
    if ext == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(f"--- Страница {idx + 1} ---\n{page_text.strip()}")
            if not pages_text:
                return "В PDF-документе не удалось извлечь печатный текст (возможно, это отсканированное изображение без текстового слоя).", False
            return "\n\n".join(pages_text), False
        except Exception as e:
            logger.exception("Ошибка парсинга PDF: %s", e)
            return f"Ошибка при чтении PDF-файла: {e}", False

    # 3. Word DOCX файл
    if ext in {".docx", ".doc"}:
        if ext == ".doc":
            return "Старый формат .doc не поддерживается напрямую. Пожалуйста, сохраните файл в современном формате .docx.", False
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Также извлекаем текст из таблиц
            for table in doc.tables:
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_lines.append(" | ".join(cells))
                if table_lines:
                    full_text.append("\n[Таблица]:\n" + "\n".join(table_lines))

            if not full_text:
                return "Документ Word не содержит текста.", False
            return "\n\n".join(full_text), False
        except Exception as e:
            logger.exception("Ошибка парсинга DOCX: %s", e)
            return f"Ошибка при чтении Word-документа: {e}", False

    # 4. Текстовые файлы и код
    if ext in TEXT_EXTENSIONS or (mime_type and ("text" in mime_type or "json" in mime_type)):
        for encoding in ["utf-8", "utf-8-sig", "cp1251", "latin-1"]:
            try:
                return file_bytes.decode(encoding), False
            except UnicodeDecodeError:
                continue
        return "Не удалось декодировать текстовый файл в поддерживаемых кодировках (utf-8, cp1251).", False

    # 5. По умолчанию пробуем раскодировать как текст
    try:
        text = file_bytes.decode("utf-8")
        return text, False
    except UnicodeDecodeError:
        return f"Формат файла '{ext}' не поддерживается для прямого текстового анализа.", False
